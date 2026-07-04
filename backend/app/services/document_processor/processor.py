import pandas as pd
import fitz  # PyMuPDF
from typing import List, Dict, Any
import os

def process_document(file_path: str) -> Dict[str, Any]:
    """
    Extract text and data from various document types.
    """
    ext = os.path.splitext(file_path)[1].lower()
    
    if ext == ".pdf":
        return process_pdf(file_path)
    elif ext in [".csv"]:
        return process_csv(file_path)
    elif ext in [".xlsx", ".xls"]:
        return process_excel(file_path)
    elif ext in [".docx"]:
        return process_docx(file_path)
    else:
        raise ValueError(f"Unsupported file type: {ext}")

def process_pdf(file_path: str) -> Dict[str, Any]:
    doc = fitz.open(file_path)
    text = ""
    for page in doc:
        text += page.get_text()
    
    return {
        "text": text,
        "type": "pdf",
        "metadata": {"pages": len(doc)}
    }

def process_csv(file_path: str) -> Dict[str, Any]:
    df = pd.read_csv(file_path)
    return {
        "text": df.to_string(),
        "type": "csv",
        "metadata": {"rows": len(df), "columns": list(df.columns)}
    }

def process_excel(file_path: str) -> Dict[str, Any]:
    df = pd.read_excel(file_path)
    return {
        "text": df.to_string(),
        "type": "excel",
        "metadata": {"rows": len(df), "columns": list(df.columns)}
    }

def process_docx(file_path: str) -> Dict[str, Any]:
    # Placeholder for docx processing if python-docx is not installed
    # To use: pip install python-docx
    try:
        import docx
        doc = docx.Document(file_path)
        text = "\n".join([para.text for para in doc.paragraphs])
        return {
            "text": text,
            "type": "docx",
            "metadata": {"paragraphs": len(doc.paragraphs)}
        }
    except ImportError:
        return {
            "text": "DOCX processing requires python-docx. Please install it.",
            "type": "docx",
            "metadata": {}
        }
